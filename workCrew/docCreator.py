import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import qn
from docx.shared import Pt
from workCrew.models import Student, Job, Day_Info 
import random
import datetime
#import arrow test

def jobAssignment(jobList, schdict, namelist, jnamelist):
    a = 0    
    b = 3
    for j in jobList:
        jnamelist.append(j.work_Name)
        schdict[j.work_Name] = namelist[a:b]
        del namelist[a:b]
        a+=3
        b+=3
        
def stringifyDictionary(unstrung, theString):
    for wele in unstrung:
        theString += (wele + "  ")    


def stringifyJobs(jnamelist, schdict):
    workString = ""
    for w in jnamelist:
        if workString != "":
            workString += "\n"            
        workString += (w + "- ")
        for wele in schdict[w]:
                workString += (wele + "  ") 
        
        
        
    return workString

def halfList(list, firstSecond):
    firstHalf = []
    secondHalf = []   
    firstHalf = list[0:((len(list)-1)/2)]
    secondHalf = list[((len(list) -1)/2): len(list)]
    
    if firstSecond == "firstHalf":
        return firstHalf
        
    elif firstSecond == "secondHalf":
        return secondHalf
  

# def lunchAssignment(eaters):
    # lunches = {}
    # for eat in eaters:
        # if len(eaters) < 15:
            # lunches[a] = eaters # comeback


def docCreator():
    weekdays = Day_Info.objects.all().order_by('order')
    
    checkDate = datetime.date.today() + datetime.timedelta(days = 1)
    
    
    weekList = []
    weekdayNumber = checkDate.weekday()
    wkcounter = 0
    while wkcounter < 7:
        weekList.append(weekdayNumber)
        if weekdayNumber == 6:
            weekdayNumber = 0
        else:
            weekdayNumber += 1
        wkcounter += 1
    weekdays3 = [weekdays[i] for i in weekList]
    
    
    for weekday in weekdays3:      
        
        
        amJobs = Job.objects.filter(day__day__contains = weekday).filter(time='am')
        beforeBreakfastJobs = Job.objects.filter(day__day__contains = weekday).filter(time='before_breakfast')
        pmJobs = Job.objects.filter(day__day__contains = weekday).filter(time='pm')
        afterDinnerJobs = Job.objects.filter(day__day__contains = weekday).filter(time='after_dinner')
        
        #jobs = Job.objects.all()
        students = Student.objects.filter(arrival_Date__lt=checkDate).filter(departure_Date__gt=checkDate)
        arrivals = Student.objects.filter(arrival_Date = checkDate)
        departures = Student.objects.filter(departure_Date = checkDate)
        strArrivals = ""
        strDepartures = ""
        for a in arrivals:
            strArrivals += (a.first_Name + " ")
            
        for d in departures:
            strDepartures += (d.first_Name + " ")
       
       
        jobnames = []
        jobnamesB = []
        
        studentnames = []  
        mealstudentNames = []
        strstudentnames = ""
        mealplace = ["a","b"]
        
        schdictAM = {}
        schdictPM = {}
        mealdict = {}
        
        for s in students:
            studentnames.append(s.first_Name)
            mealstudentNames.append(s.first_Name)
            strstudentnames += (s.first_Name + " ")
        
        
        random.shuffle(studentnames)
        random.shuffle(mealstudentNames)
         
        
        
        jobAssignment(beforeBreakfastJobs,schdictAM, studentnames,jobnames)    
        jobAssignment(amJobs, schdictAM, studentnames, jobnames)        
        
        jobAssignment(pmJobs, schdictPM, studentnames, jobnamesB)
        jobAssignment(afterDinnerJobs, schdictPM, studentnames, jobnamesB)
        
        
        
        
     

        
        
        workAM = stringifyJobs(jobnames, schdictAM)       
        workPM = stringifyJobs(jobnamesB, schdictPM)       
        
        
        
        document = docx.Document()#creating document
        
        
        #assign day
        pweekday = checkDate.strftime('%A' ) + " " + checkDate.strftime('%B' ) +  " " + str(checkDate.day) 
        
        checkDate += datetime.timedelta(days=1) #increment day
        
        #pweekday = str(weekdays) + str(weekdays3)
        day = document.add_paragraph(pweekday)
        day_format = day.paragraph_format
        day_format.space_after = Pt(18)
        day_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        #breakfast infor
        breakfastTime = document.add_paragraph("8:00am Breakfast -(everyone expected)" )
        breakfastTime_format = breakfastTime.paragraph_format
        breakfastTime_format.space_after = Pt(18)
        breakfastTime_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        locationB = document.add_paragraph("At Chalet Bellevue with " + weekday.breakfast)
        locationB_format = locationB.paragraph_format
        locationB_format.space_after = Pt(18)
        locationB_format.alignment = WD_ALIGN_PARAGRAPH.CENTER        
        
        
        tableB = document.add_table(rows=1, cols=1)
        tableB.borders={'All':Pt(.2)}
        row1B = tableB.rows[0]
        cell1B = row1B.cells[0]
        if strDepartures:
            cell1B.text = strstudentnames + strDepartures + "\nMeal Number:  " +  str(len(mealstudentNames) + len(departures))
        else:
            cell1B.text = strstudentnames + "\nMeal Number:  " +  str(len(mealstudentNames))
            
        if weekday.day == "Monday":
            prayerTime = document.add_paragraph("8:45am Prayer Meeting in the Lounge" )
            prayerTime_format = prayerTime.paragraph_format
            prayerTime_format.space_after = Pt(18)
            prayerTime_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        if weekday.lecture != "none" and weekday.day == "Friday":
            lectureTime = document.add_paragraph("9:15am Lecture in Farel House" )
            lectureTime_format = lectureTime.paragraph_format
            lectureTime_format.space_after = Pt(18)
            lectureTime_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_arrivals = document.add_paragraph("Arrivals: " + strArrivals)
        p_arrivals_format = p_arrivals.paragraph_format
        p_arrivals_format.space_before = Pt(18)
        p_arrivals_format.space_after = Pt(0)
        p_arrivals_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_departures = document.add_paragraph("Departures: " + strDepartures)
        p_departures_format = p_departures.paragraph_format
        p_departures_format.space_after = Pt(18)
        p_departures_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        
        if (len(mealstudentNames) + len(arrivals)) > 14:
            if weekday.Lunch_A != 'none':
                lunchTime = document.add_paragraph("1:00 Lunch-(everyone expected)")    
                lunchTime_format = lunchTime.paragraph_format
                lunchTime_format.space_after = Pt(18)
                lunchTime_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                locationLa = document.add_paragraph("At Chalet " + weekday.Lunch_A)#plus the location  
                locationLa_format = locationLa.paragraph_format
                locationLa_format.space_after = Pt(18)
                locationLa_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                tableLa = document.add_table(rows=1, cols=1)
                row1La = tableLa.rows[0]
                cell1La = row1La.cells[0]
                if strArrivals:
                    cell1La.text = ' '.join(halfList(mealstudentNames, "firstHalf")) + ("("+ strArrivals + "?" + ")") + "\nMeal Number:  " +  str(len(mealstudentNames) + len(arrivals))
                else:
                    cell1La.text = strstudentnames + "\nMeal Number:  " +  str(len(mealstudentNames))
                
                locationLb = document.add_paragraph("At Chalet " + weekday.Lunch_B)#plus the location         
                locationLb_format = locationLb.paragraph_format
                locationLb_format.space_after = Pt(18)
                locationLb_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                tableLb = document.add_table(rows=1, cols=1)
                row1Lb = tableLb.rows[0]
                cell1Lb = row1Lb.cells[0]                
                cell1Lb.text = ' '.join(halfList(mealstudentNames, "secondHalf"))
            
            else:
                locationLa = document.add_paragraph("Packed lunches are in the student fridge")
                locationLa_format = locationLa.paragraph_format
                locationLa_format.space_after = Pt(18)
                locationLa_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            
            
            

        else:
            if weekday.Lunch_A != 'none':
                lunchTime = document.add_paragraph("1:00 Lunch-(everyone expected)")    
                lunchTime_format = lunchTime.paragraph_format
                lunchTime_format.space_after = Pt(18)
                lunchTime_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                locationLa = document.add_paragraph("At Chalet " + weekday.Lunch_A)#plus the location  
                locationLa_format = locationLa.paragraph_format
                locationLa_format.space_after = Pt(18)
                locationLa_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                tableLa = document.add_table(rows=1, cols=1)
                row1La = tableLa.rows[0]
                cell1La = row1La.cells[0]
                if strArrivals:
                    cell1La.text = strstudentnames + ("("+ strArrivals + "?" + ")") + "\nMeal Number:  " +  str(len(mealstudentNames) + len(arrivals))
                else:
                    cell1La.text = strstudentnames + "\nMeal Number:  " +  str(len(mealstudentNames))
            
            else:
                locationLa = document.add_paragraph("Packed lunches are in the student fridge")
                locationLa_format = locationLa.paragraph_format
                locationLa_format.space_after = Pt(18)
                locationLa_format.alignment = WD_ALIGN_PARAGRAPH.CENTER        
        
        
        if weekday.lecture != "none" and weekday.day == "Wednesday":
            lectureTime = document.add_paragraph("2:00pm Lecture in Farel House" )
            lectureTime_format = lectureTime.paragraph_format
            lectureTime_format.space_after = Pt(18)
            lectureTime_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        dinnerTime = document.add_paragraph("6:30 Dinner -(everyone expected)")    
        dinnerTime_format = dinnerTime.paragraph_format
        dinnerTime_format.space_before = Pt(40)
        dinnerTime_format.space_after = Pt(18)
        dinnerTime_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        locationD = document.add_paragraph("At Chalet Bellevue with " + weekday.dinner)
        locationD_format = locationD.paragraph_format
        locationD_format.space_after = Pt(18)
        locationD_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tableD = document.add_table(rows=1, cols=1)
        row1D = tableD.rows[0]
        cell1D = row1D.cells[0]
        if strArrivals:
            cell1D.text = strstudentnames + ("("+ strArrivals + "?" + ")") + "\nMeal Number:  " +  str(len(mealstudentNames) + len(arrivals))
        else:
            cell1D.text = strstudentnames + "\nMeal Number:  " + str(len(mealstudentNames))
        
        workPara = document.add_paragraph("Work Assignments: if your name is not down for a work crew, you have a study day:")    
        workPara_format = workPara.paragraph_format
        workPara_format.space_before = Pt(20)
        workPara_format.space_after = Pt(6)
        workPara_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tablew = document.add_table(rows=1, cols=2)
        cell1w = tablew.rows[0].cells[0]
        cell1w.text = ("am\n" + "\n 9:30am: \n \n" + workAM)
        cell2w = tablew.rows[0].cells[1]
        cell2w.text = ("pm\n \n 3:00pm:\n \n" + workPM )
        document.save('test/' + pweekday +'.docx')
        